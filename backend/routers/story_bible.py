"""
Phase 2 — P2-06 Story Bible Generator
Synthesises all manuscript data into a structured story bible via Qwen.
Stored in story_bibles table. Exportable as DOCX.
Background task with polling-compatible status approach.
"""
import asyncio
import json
import logging
import uuid
from dataclasses import dataclass
from datetime import datetime

from fastapi import APIRouter, Depends, HTTPException, Request
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session

from config import settings
from database import get_db
from exceptions import AIServiceUnavailableError
from middleware.rate_limit import limiter, get_user_id
from middleware.concurrency import bg_ai_semaphore
from models import Story, ChapterSummary, CharacterProfile, Character, StoryNote, NoteCard, GenreProfile, StoryBible
from schemas import StoryBibleOut, StoryBibleJobResponse
from routers.auth import get_current_user, User
from services.ai_service import audit_section_provenance, count_tokens, generate_story_bible_section

logger = logging.getLogger(__name__)

router = APIRouter(tags=["story-bible"])

_BIBLE_SECTIONS = ["characters", "locations", "timeline", "world_rules", "themes"]

# In-process guard: prevent concurrent bible generation for the same story
_generating: set[str] = set()


# ── Per-section generation outcomes ───────────────────────────────────────────
#
# Why this exists: the pipeline used to record failure only in the *content*
# (an "[AI temporarily unavailable …]" string written where the section text
# belongs) and never in any variable the persistence step could read. With no
# per-section outcome to consult, the terminal status could only ever be
# hard-coded 'completed' — which is how a bible with five failed sections came
# to report success. Failure paths SB-F16 and SB-F17 in
# docs/issues-and-bugs/story-bible-failure-path-audit.md.
#
# Two of the failure classes below raise nothing at all, so an exception-driven
# record would miss them: EMPTY (SB-F12) and TRUNCATED (SB-F13). Outcomes are
# therefore classified from the *response*, not only from raised exceptions.
#
# These outcomes drive everything the pipeline persists about failure:
# derive_status() turns them into completed / partial / failed, and
# failed_section_payload() into the failed_sections rows. Failure text is no
# longer written into content_json at all — a section that did not produce
# genuine content is simply absent from it.

FAIL_UNAVAILABLE = "unavailable"   # AI service unreachable / overloaded
FAIL_ERROR       = "error"         # any other exception during generation
FAIL_EMPTY       = "empty"         # returned nothing, or only whitespace
FAIL_TRUNCATED   = "truncated"     # stopped at max_tokens — cut off mid-section

# Generation lifecycle. 'running' is written by the POST handler; the three
# terminal values below are derived from the section outcomes, never from the
# fact that the loop finished.
STATUS_RUNNING   = "running"
STATUS_COMPLETED = "completed"     # every section produced genuine content
STATUS_PARTIAL   = "partial"       # some sections did, some did not
STATUS_FAILED    = "failed"        # none did, or the pipeline never got that far


@dataclass(frozen=True)
class SectionOutcome:
    """One section's generation result. ``reason`` is author-safe text — it
    never carries manuscript content or an exception repr."""
    section: str
    ok:      bool
    failure: str | None = None     # one of the FAIL_* constants when not ok
    reason:  str = ""


def classify_section_result(
    section: str,
    text: str | None,
    finish_reason: str | None,
) -> SectionOutcome:
    """
    Decide whether a returned section actually holds genuine content.

    Called only on the path where the AI call returned *without* raising —
    because an exception-free response is not the same as a usable one. Order
    matters: an empty response that also hit the token cap is empty, not
    truncated, so emptiness is checked first.
    """
    if text is None or not text.strip():
        return SectionOutcome(
            section, False, FAIL_EMPTY,
            "The AI returned nothing for this section.",
        )
    if finish_reason == "length":
        return SectionOutcome(
            section, False, FAIL_TRUNCATED,
            "This section was cut off before it finished.",
        )
    return SectionOutcome(section, True)


def failed_section_payload(outcome: SectionOutcome) -> dict:
    """
    The persisted, author-facing record of one failed section.

    Deliberately minimal: which section, which failure class, and what to tell
    the author. Exception text, stack traces, timestamps and model diagnostics
    stay in the logs — this is an API contract the frontend will render, not a
    debugging channel.
    """
    return {
        "section": outcome.section,
        "failure": outcome.failure,
        "reason":  outcome.reason,
    }


def derive_status(outcomes: list[SectionOutcome]) -> str:
    """
    Terminal status for a finished generation run, derived from what the
    sections actually produced.

    The rule this replaces was ``status = "completed"`` written unconditionally
    once the loop ended, which is how a bible whose five sections were all
    error placeholders came to report success (SB-F20).

    A section that failed contributes no content at all — nothing but genuinely
    generated text is stored — so a bible with any failure can never read
    'completed'.

    No outcomes at all means the loop produced nothing to judge, which is a
    failure, not a success.
    """
    if not outcomes:
        return STATUS_FAILED
    generated = sum(1 for o in outcomes if o.ok)
    if generated == len(outcomes):
        return STATUS_COMPLETED
    if generated == 0:
        return STATUS_FAILED
    return STATUS_PARTIAL


def outcomes_from_persisted_state(content: dict, failed_sections: list | None) -> list[SectionOutcome]:
    """
    Reconstruct the per-section outcomes of a stored bible.

    Needed when only one section is regenerated: the status must be recomputed
    from the whole bible, not from the single section that just ran. A section
    present in content_json produced genuine content (subtask 5 guarantees
    nothing else is stored there); one named in failed_sections did not.

    A section in neither is treated as failed — that is a bible written before
    failed_sections existed, or one whose row was edited outside the app, and
    'we have no content for this' is the honest reading either way.
    """
    failures = {f.get("section"): f for f in (failed_sections or []) if isinstance(f, dict)}
    outcomes: list[SectionOutcome] = []
    for name in _BIBLE_SECTIONS:
        if content.get(name):
            outcomes.append(SectionOutcome(name, True))
        else:
            entry = failures.get(name, {})
            outcomes.append(SectionOutcome(
                name, False,
                entry.get("failure") or FAIL_ERROR,
                entry.get("reason") or "This section has not been generated yet.",
            ))
    return outcomes


def _get_owned_story(story_id: str, user_id: str, db: Session) -> Story:
    story = db.query(Story).filter(
        Story.story_id == story_id,
        Story.user_id  == user_id,
    ).first()
    if not story:
        raise HTTPException(status_code=404, detail="Story not found")
    return story


# ── Context assembly ──────────────────────────────────────────────────────────
#
# The bible used to be built from chapter summaries truncated to 300 characters
# each — a fixed cap applied regardless of how many chapters existed, so the
# model was shown a digest and asked for a chronicle, and filled the gaps by
# inventing. That is the grounding defect (Phase 2 Issue 8; SB-F03/SB-F04).
#
# Three things replace it:
#   1. a budget derived from the model's actual context window and measured in
#      real Qwen tokens, so "as much as fits" is a fact rather than a guess;
#   2. a provenance tag on every entry, so the model can cite what it used;
#   3. visible degradation — when a manuscript cannot fit, the context says so
#      instead of silently dropping the end of the book.

# Must match the max_tokens used by generate_story_bible_section.
_BIBLE_COMPLETION_TOKENS = 1500
# System prompt, per-section instruction, framing and a safety margin.
_BIBLE_PROMPT_OVERHEAD = 900
# A chapter summary below this is barely worth including; it is also the old
# fixed cap, so an included chapter is never worse off than before.
_MIN_SUMMARY_CHARS = 300


def bible_context_token_budget() -> int:
    """Input tokens available for manuscript context, from the serving window."""
    max_ctx = getattr(settings, "max_model_len", 8192) or 8192
    return max(1000, max_ctx - _BIBLE_COMPLETION_TOKENS - _BIBLE_PROMPT_OVERHEAD)


def fair_share(sizes: list[int], budget: int) -> list[int]:
    """
    Max–min fair allocation of `budget` across items of the given sizes.

    Small entries get everything they need; whatever is left is split evenly
    among the large ones. Prevents one enormous chapter summary from consuming
    the window and starving the other thirty — which is what a naive
    first-come fill does, and why the timeline was the worst-affected section.
    """
    alloc = [0] * len(sizes)
    remaining, left = budget, len(sizes)
    for i in sorted(range(len(sizes)), key=lambda i: sizes[i]):
        share = remaining // left if left else 0
        alloc[i] = min(sizes[i], share)
        remaining -= alloc[i]
        left -= 1
    return alloc


def _summary_entry(s: ChapterSummary, body_chars: int) -> str:
    """One chapter summary, tagged with its provenance so it can be cited."""
    body = str(s.raw_summary or "")
    if body_chars and len(body) > body_chars:
        body = body[:body_chars].rstrip() + "…"
    return (
        f"[Ch {s.chapter_number}] "
        f"events={s.key_events}, "
        f"characters={s.characters_present}, "
        f"locations={s.locations}, "
        f"tone={s.emotional_tone}, "
        f"summary={body}"
    )


def _build_full_context(story_id: str, db: Session) -> str:
    """
    Assemble manuscript data into a single, provenance-tagged context string
    sized to the model's context window.

    Every entry carries a tag — ``[Ch 7]``, ``[Character: Devika Rao]`` — so the
    generated bible can cite its source and a reader can check it. When the
    manuscript does not fit, the context states plainly what was shortened or
    left out; the model is told not to describe what it was not shown.
    """
    budget = bible_context_token_budget()
    notes_lines: list[str] = []      # what the model is told about its own limits

    # ── Chapter summaries — the primary evidence, budgeted first ──────────────
    summaries = (
        db.query(ChapterSummary)
        .filter(ChapterSummary.story_id == story_id)
        .order_by(ChapterSummary.chapter_number)
        .all()
    )

    lines: list[str] = []
    included_chapters: list[int] = []
    omitted_chapters: list[int] = []
    shortened = 0

    if summaries:
        # Reserve most of the budget for summaries; the rest of the context
        # (profiles, notes, cards) shares what remains.
        summary_token_budget = int(budget * 0.75)
        # Convert to characters using the measured ratio for this manuscript,
        # rather than assuming one — prose and metadata tokenise differently.
        raw_bodies = [str(s.raw_summary or "") for s in summaries]
        sample = "\n".join(raw_bodies) or " "
        chars_per_token = max(1.0, len(sample) / max(1, count_tokens(sample)))
        char_budget = int(summary_token_budget * chars_per_token)

        # Fixed cost of each entry's metadata, so the allocation covers bodies only.
        overheads = [len(_summary_entry(s, 0)) - len(str(s.raw_summary or "")) for s in summaries]
        avg_overhead = sum(overheads) // len(overheads)

        # How many chapters can carry at least a useful summary? Past that point
        # spreading the budget thinner produces stubs the model would still cite,
        # so chapters are dropped instead — but never ALL of them: an allocation
        # that leaves no manuscript at all is worse than a partial view, because
        # the model then has nothing to ground on and the section is pure
        # invention. (Measured: a 200-chapter manuscript hit exactly that.)
        capacity = char_budget // max(1, _MIN_SUMMARY_CHARS + avg_overhead)

        if capacity < len(summaries):
            # Sample evenly across the whole manuscript rather than taking the
            # first N. A contiguous prefix would tell the model nothing about
            # how the story ends, and a story bible built from act one is a
            # different kind of wrong. The omission notice names what is missing.
            step = len(summaries) / max(1, capacity)
            keep = {int(i * step) for i in range(max(1, capacity))}
            keep.add(len(summaries) - 1)   # always keep the ending — a bible
                                           # that stops before the climax is a
                                           # bible of a different book
            chosen = [(i, s) for i, s in enumerate(summaries) if i in keep]
            omitted_chapters = [s.chapter_number for i, s in enumerate(summaries) if i not in keep]
        else:
            chosen = list(enumerate(summaries))

        chosen_bodies = [raw_bodies[i] for i, _ in chosen]
        chosen_overhead = sum(overheads[i] for i, _ in chosen)
        body_budget = max(0, char_budget - chosen_overhead)
        allocations = fair_share([len(b) for b in chosen_bodies], body_budget)

        for (_, s), body, alloc in zip(chosen, chosen_bodies, allocations):
            included_chapters.append(s.chapter_number)
            if alloc < len(body):
                shortened += 1
            lines.append(_summary_entry(s, alloc))

        if lines:
            lines.insert(0, "=== CHAPTER SUMMARIES ===")

    used = count_tokens("\n".join(lines))
    remaining = max(0, budget - used)

    def _add_section(header: str, entries: list[str]) -> int:
        """Append what fits of `entries`; return how many were left out."""
        nonlocal remaining
        if not entries:
            return 0
        header_cost = count_tokens(header)
        if header_cost >= remaining:
            return len(entries)
        lines.append(header)
        remaining -= header_cost
        dropped = 0
        for entry in entries:
            cost = count_tokens(entry)
            if cost > remaining:
                dropped += 1
                continue
            lines.append(entry)
            remaining -= cost
        return dropped

    # ── Character profiles ────────────────────────────────────────────────────
    chars = (
        db.query(Character, CharacterProfile)
        .outerjoin(CharacterProfile, Character.character_id == CharacterProfile.character_id)
        .filter(Character.story_id == story_id)
        .all()
    )
    char_entries = []
    for c, p in chars:
        if p:
            char_entries.append(
                f"[Character: {c.name}] ({c.role}): "
                f"appearance={(p.appearance or '')[:400]}, "
                f"personality={(p.personality or '')[:400]}, "
                f"goals={(p.goals or '')[:300]}, "
                f"backstory={(p.backstory or '')[:500]}"
            )
        else:
            char_entries.append(f"[Character: {c.name}] ({c.role}): no profile recorded")
    dropped_chars = _add_section("\n=== CHARACTER PROFILES ===", char_entries)

    # ── Story notes ───────────────────────────────────────────────────────────
    notes = db.query(StoryNote).filter(StoryNote.story_id == story_id).all()
    dropped_notes = _add_section(
        "\n=== STORY NOTES (author's own notes, not manuscript text) ===",
        [f"[Note: {n.title}] {(n.content or '')[:600]}" for n in notes],
    )

    # ── Note cards ────────────────────────────────────────────────────────────
    cards = db.query(NoteCard).filter(NoteCard.story_id == story_id).all()
    dropped_cards = _add_section(
        "\n=== NOTE CARDS (author's own notes, not manuscript text) ===",
        [f"[Card: {c.card_type} — {c.title}] {(c.content or '')[:400]}" for c in cards],
    )

    # ── Genre ─────────────────────────────────────────────────────────────────
    genre = db.query(GenreProfile).filter(GenreProfile.story_id == story_id).first()
    if genre and remaining > 50:
        lines.append(f"\n=== GENRE ===\nGenre: {genre.genre}, Sub-genre: {genre.sub_genre}, Tone: {genre.tone}")

    # ── Tell the model what it is NOT seeing ──────────────────────────────────
    # Silent truncation is what let the model believe it had the whole book.
    if omitted_chapters:
        shown = f"{len(included_chapters)} of {len(summaries)}"
        notes_lines.append(
            f"Only {shown} chapters are included here. Chapters not shown: "
            f"{_compact_ranges(omitted_chapters)}."
        )
    if shortened:
        notes_lines.append(f"{shortened} chapter summaries were shortened to fit and end with '…'.")
    for count, what in ((dropped_chars, "character profiles"),
                        (dropped_notes, "story notes"),
                        (dropped_cards, "note cards")):
        if count:
            notes_lines.append(f"{count} {what} were not included.")
    if notes_lines:
        lines.append(
            "\n=== WHAT YOU ARE NOT SEEING ===\n"
            + "\n".join(notes_lines)
            + "\nDo not describe anything from material that is not shown above. "
              "If the bible would need it, say so instead of inferring it."
        )

    context = "\n".join(lines)
    logger.info(
        "[story_bible] context for %s: %d tokens (budget %d), %d/%d chapters, "
        "%d shortened, dropped %d profiles / %d notes / %d cards",
        story_id[:8], count_tokens(context), budget,
        len(included_chapters), len(summaries), shortened,
        dropped_chars, dropped_notes, dropped_cards,
    )
    return context


def _log_provenance(story_id: str, section: str, text: str) -> None:
    """
    Record how well a generated section cited its sources.

    Diagnostics only — never shown to the author, never used to reject content.
    Counts and ratios only, so no manuscript text reaches the log.
    """
    stats = audit_section_provenance(text)
    if not stats["entries"]:
        return
    log = logger.warning if (stats["cited_ratio"] or 0) < 0.5 else logger.info
    log("[story_bible] provenance %s/%s: %d/%d entries cited (%.0f%%), %d marked not-established",
        story_id[:8], section, stats["cited"], stats["entries"],
        100 * (stats["cited_ratio"] or 0), stats["not_established"])


def _compact_ranges(numbers: list[int]) -> str:
    """[1,2,3,7,9,10] → '1-3, 7, 9-10' — keeps the omission note short."""
    if not numbers:
        return ""
    ordered = sorted(set(numbers))
    parts, start, prev = [], ordered[0], ordered[0]
    for n in ordered[1:] + [None]:
        if n is not None and n == prev + 1:
            prev = n
            continue
        parts.append(str(start) if start == prev else f"{start}-{prev}")
        if n is not None:
            start = prev = n
    return ", ".join(parts)


async def _generate_bible_pipeline(story_id: str, bible_id: str, bump_version: bool) -> None:
    """
    Background task: generate all 5 sections into the pre-created ``bible_id`` row
    (created with status='running' by the POST handler), then set the terminal
    status from what the sections actually produced — completed / partial /
    failed, see derive_status(). On a pipeline-level failure the row is marked
    'failed' so the UI can offer a retry.
    Runs under bg_ai_semaphore to prevent vLLM queue flooding.
    """
    from database import SessionLocal
    db = SessionLocal()
    try:
        _generating.add(story_id)
        context = _build_full_context(story_id, db)

        content: dict[str, str] = {}
        outcomes: list[SectionOutcome] = []
        async with bg_ai_semaphore():
            for section in _BIBLE_SECTIONS:
                # Only a section that produced genuine content is stored. A
                # failed one is ABSENT from content_json and described by
                # failed_sections — the status and that list are the canonical
                # record of failure. Nothing this code generates about its own
                # failure is ever written where a section belongs (SB-F16).
                #
                # Truncated sections are dropped too, deliberately: the prose is
                # real but stops mid-sentence, and a fragment kept here would be
                # rendered and exported as if it were a finished section. The
                # bible stores usable sections only; intentionally preserving an
                # unfinished generation is the pinning feature's job.
                try:
                    text, finish_reason = await generate_story_bible_section(section=section, context=context)
                    outcome = classify_section_result(section, text, finish_reason)
                    if outcome.ok:
                        content[section] = text
                        _log_provenance(story_id, section, text)
                    else:
                        logger.warning("[story_bible] section %s not usable (story=%s): %s",
                                       section, story_id[:8], outcome.failure)
                    outcomes.append(outcome)
                except AIServiceUnavailableError as exc:
                    logger.warning("[story_bible] AI unavailable for section %s (story=%s): %s", section, story_id[:8], exc)
                    outcomes.append(SectionOutcome(
                        section, False, FAIL_UNAVAILABLE,
                        "The AI service was unavailable while writing this section.",
                    ))
                except Exception as exc:
                    # The exception text stays here, in the log. It is not
                    # author-safe and must never reach content_json, the API
                    # payload or the DOCX export.
                    logger.warning("[story_bible] section %s failed (story=%s): %s: %s",
                                   section, story_id[:8], type(exc).__name__, exc)
                    outcomes.append(SectionOutcome(
                        section, False, FAIL_ERROR,
                        "This section could not be generated.",
                    ))

        # Section keys and failure classes only — never section text.
        generated = [o.section for o in outcomes if o.ok]
        failed    = [o for o in outcomes if not o.ok]
        if failed:
            logger.warning(
                "[story_bible] %s: %d/%d sections generated; failed — %s",
                story_id[:8], len(generated), len(outcomes),
                ", ".join(f"{o.section}={o.failure}" for o in failed),
            )
        else:
            logger.info("[story_bible] %s: %d/%d sections generated",
                        story_id[:8], len(generated), len(outcomes))

        content_json = json.dumps(content, ensure_ascii=False)
        status = derive_status(outcomes)

        bible = db.query(StoryBible).filter(StoryBible.bible_id == bible_id).first()
        if bible:
            bible.content_json = content_json
            bible.status       = status
            # Always assigned, never appended: a successful regeneration must
            # clear the previous run's failures, or a repaired bible would keep
            # reporting sections that are now fine.
            bible.failed_sections = [failed_section_payload(o) for o in failed]
            if bump_version:
                bible.version  = (bible.version or 1) + 1
            bible.updated_at   = datetime.utcnow()
            db.commit()
            logger.info("[story_bible] %s v%d for %s (%d/%d sections)",
                        status, bible.version, story_id[:8], len(generated), len(outcomes))
    except Exception as exc:
        logger.error("[story_bible] generation failed for %s: %s", story_id[:8], exc)
        try:
            # Roll back FIRST. If the exception we are handling came from the
            # commit above, this session's transaction has already failed and
            # SQLAlchemy refuses any further query on it — the recovery query
            # would raise PendingRollbackError, the inner handler would swallow
            # it, and the row would be left in 'running' for ever, un-retryable
            # because the duplicate guard treats 'running' as work in progress.
            # Only a restart cleared it. Rolling back first is what makes
            # 'failed' reachable on the commit path at all (SB-F23).
            db.rollback()
            bible = db.query(StoryBible).filter(StoryBible.bible_id == bible_id).first()
            if bible:
                bible.status     = STATUS_FAILED
                # A pipeline-level failure carries no per-section detail, so the
                # previous run's list must not be left behind describing failures
                # that are not why this run failed.
                bible.failed_sections = []
                bible.updated_at = datetime.utcnow()
                db.commit()
                logger.info("[story_bible] marked failed for %s", story_id[:8])
        except Exception as recovery_exc:
            logger.error("[story_bible] could not mark %s failed: %s: %s — row left in '%s', "
                         "startup orphan recovery will clear it",
                         story_id[:8], type(recovery_exc).__name__, recovery_exc, STATUS_RUNNING)
            db.rollback()
    finally:
        _generating.discard(story_id)
        db.close()


async def _regenerate_section_pipeline(story_id: str, bible_id: str, section: str) -> None:
    """
    Background task: regenerate ONE section into an existing bible, merge it,
    and recompute the terminal status across the whole bible.

    Differs from the full pipeline in three ways, each deliberate:

      * the version is never bumped — repairing a section completes an existing
        bible rather than creating a new canonical one;
      * the status is recomputed from the merged state of all five sections,
        not from this one run, so fixing the last failure yields 'completed'
        and fixing one of several leaves 'partial';
      * a pipeline-level failure does NOT mark the whole bible failed. Four good
        sections are still four good sections; the status is restored from what
        is actually persisted, so a failed repair attempt cannot destroy the
        author's standing result.
    """
    from database import SessionLocal
    db = SessionLocal()
    try:
        _generating.add(story_id)
        context = _build_full_context(story_id, db)

        async with bg_ai_semaphore():
            try:
                text, finish_reason = await generate_story_bible_section(section=section, context=context)
                outcome = classify_section_result(section, text, finish_reason)
                if outcome.ok:
                    _log_provenance(story_id, section, text)
            except AIServiceUnavailableError as exc:
                logger.warning("[story_bible] AI unavailable regenerating %s (story=%s): %s", section, story_id[:8], exc)
                outcome = SectionOutcome(section, False, FAIL_UNAVAILABLE,
                                         "The AI service was unavailable while writing this section.")
            except Exception as exc:
                logger.warning("[story_bible] regenerating %s failed (story=%s): %s: %s",
                               section, story_id[:8], type(exc).__name__, exc)
                outcome = SectionOutcome(section, False, FAIL_ERROR,
                                         "This section could not be generated.")

        bible = db.query(StoryBible).filter(StoryBible.bible_id == bible_id).first()
        if bible:
            try:
                content = json.loads(bible.content_json or "{}")
            except (json.JSONDecodeError, TypeError):
                content = {}

            # Merge: a usable section replaces what was there; a failed one
            # leaves no content behind (subtask 5 — nothing but genuine content
            # is ever stored).
            if outcome.ok:
                content[section] = text
            else:
                content.pop(section, None)

            others = [f for f in (bible.failed_sections or [])
                      if isinstance(f, dict) and f.get("section") != section]
            failed_now = others if outcome.ok else others + [failed_section_payload(outcome)]

            bible.content_json    = json.dumps(content, ensure_ascii=False)
            bible.failed_sections = failed_now
            bible.status          = derive_status(outcomes_from_persisted_state(content, failed_now))
            bible.updated_at      = datetime.utcnow()
            db.commit()
            logger.info("[story_bible] section %s regenerated for %s → %s (%d still failing)",
                        section, story_id[:8], bible.status, len(failed_now))
    except Exception as exc:
        logger.error("[story_bible] section regeneration failed for %s/%s: %s", story_id[:8], section, exc)
        try:
            db.rollback()   # same ordering rule as the full pipeline (SB-F23)
            bible = db.query(StoryBible).filter(StoryBible.bible_id == bible_id).first()
            if bible:
                # Restore the status the persisted content actually supports,
                # rather than declaring the whole bible failed over one section.
                try:
                    stored = json.loads(bible.content_json or "{}")
                except (json.JSONDecodeError, TypeError):
                    stored = {}
                bible.status     = derive_status(outcomes_from_persisted_state(stored, bible.failed_sections))
                bible.updated_at = datetime.utcnow()
                db.commit()
                logger.info("[story_bible] %s restored to %s after failed section repair",
                            story_id[:8], bible.status)
        except Exception as recovery_exc:
            logger.error("[story_bible] could not restore %s after failed repair: %s: %s — row left in '%s', "
                         "startup orphan recovery will clear it",
                         story_id[:8], type(recovery_exc).__name__, recovery_exc, STATUS_RUNNING)
            db.rollback()
    finally:
        _generating.discard(story_id)
        db.close()


@router.post("/{story_id}/story-bible", response_model=StoryBibleJobResponse)
@limiter.limit(settings.rate_limit_background_ai, key_func=get_user_id)
async def create_or_regenerate_bible(
    request: Request,
    story_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """
    Trigger story bible generation (or regeneration if one already exists).
    Returns a job_id immediately. Poll GET /story-bible until the bible appears.
    Background task: Qwen synthesises characters, locations, timeline, world rules, themes.
    """
    _get_owned_story(story_id, current_user.user_id, db)

    existing = (
        db.query(StoryBible)
        .filter(StoryBible.story_id == story_id)
        .order_by(StoryBible.version.desc())
        .first()
    )

    # ── Duplicate prevention (persisted + in-memory) ──────────────────────────
    # A row in 'running' state OR the in-memory guard means a generation is
    # already underway — return it instead of starting a second one. This
    # survives the page reload/navigation that the previous flow could not.
    if (existing and existing.status == "running") or story_id in _generating:
        return StoryBibleJobResponse(
            job_id="in-progress",
            status="already_generating",
            bible_id=existing.bible_id if existing else None,
        )

    summaries_count = (
        db.query(ChapterSummary)
        .filter(ChapterSummary.story_id == story_id)
        .count()
    )
    if summaries_count == 0:
        raise HTTPException(
            status_code=422,
            detail="No indexed chapters found. Index at least one chapter first.",
        )

    # Persist the 'running' state immediately so GET reflects in-progress even
    # before the first section is produced (and across reloads/navigation).
    bump_version = bool(existing and existing.status == "completed")
    if existing:
        existing.status     = "running"
        existing.updated_at = datetime.utcnow()
        db.commit()
        bible_id = existing.bible_id
    else:
        bible = StoryBible(
            story_id=story_id,
            user_id=current_user.user_id,
            content_json="{}",
            status="running",
        )
        db.add(bible)
        db.commit()
        db.refresh(bible)
        bible_id = bible.bible_id

    job_id = str(uuid.uuid4())
    asyncio.create_task(
        _generate_bible_pipeline(story_id, bible_id, bump_version)
    )

    return StoryBibleJobResponse(
        job_id=job_id,
        status="processing",
        bible_id=bible_id,
    )


@router.post("/{story_id}/story-bible/sections/{section}", response_model=StoryBibleJobResponse)
@limiter.limit(settings.rate_limit_background_ai, key_func=get_user_id)
async def regenerate_bible_section(
    request: Request,
    story_id: str,
    section: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """
    Regenerate a single story bible section.

    Exists so an author whose bible came back 'partial' can repair the two
    sections that failed instead of paying for all five again. Returns
    immediately; poll GET /story-bible, which reports 'running' until the
    section lands and then the recomputed status.

    The whole-bible generation guard applies — one generation at a time per
    story, whether full or targeted.
    """
    _get_owned_story(story_id, current_user.user_id, db)

    if section not in _BIBLE_SECTIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Unknown section. Expected one of: {', '.join(_BIBLE_SECTIONS)}.",
        )

    bible = (
        db.query(StoryBible)
        .filter(StoryBible.story_id == story_id)
        .order_by(StoryBible.version.desc())
        .first()
    )
    if not bible:
        raise HTTPException(
            status_code=404,
            detail="No story bible found. Generate one first.",
        )

    if bible.status == STATUS_RUNNING or story_id in _generating:
        return StoryBibleJobResponse(
            job_id="in-progress",
            status="already_generating",
            bible_id=bible.bible_id,
        )

    bible.status     = STATUS_RUNNING
    bible.updated_at = datetime.utcnow()
    db.commit()

    job_id = str(uuid.uuid4())
    asyncio.create_task(
        _regenerate_section_pipeline(story_id, bible.bible_id, section)
    )

    return StoryBibleJobResponse(
        job_id=job_id,
        status="processing",
        bible_id=bible.bible_id,
    )


@router.get("/{story_id}/story-bible", response_model=StoryBibleOut)
def get_story_bible(
    story_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Retrieve the latest generated story bible."""
    _get_owned_story(story_id, current_user.user_id, db)

    bible = (
        db.query(StoryBible)
        .filter(StoryBible.story_id == story_id)
        .order_by(StoryBible.version.desc())
        .first()
    )
    if not bible:
        raise HTTPException(
            status_code=404,
            detail="No story bible found. Generate one first via POST /story-bible.",
        )
    return bible


@router.get("/{story_id}/story-bible/export")
def export_story_bible_docx(
    story_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Export the story bible as a formatted DOCX document."""
    import io
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    story = _get_owned_story(story_id, current_user.user_id, db)

    bible = (
        db.query(StoryBible)
        .filter(StoryBible.story_id == story_id)
        .order_by(StoryBible.version.desc())
        .first()
    )
    if not bible:
        raise HTTPException(status_code=404, detail="No story bible found.")

    try:
        content = json.loads(bible.content_json)
    except (json.JSONDecodeError, TypeError):
        content = {}

    doc = DocxDocument()

    # Title page
    title_para = doc.add_heading(f"{story.title} — Story Bible", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph(f"Version {bible.version}  |  Generated by NarratIQ AI")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    section_titles = {
        "characters":  "Characters",
        "locations":   "Locations",
        "timeline":    "Timeline",
        "world_rules": "World Rules",
        "themes":      "Themes & Motifs",
    }

    for key in _BIBLE_SECTIONS:
        text = content.get(key, "")
        if not text:
            continue

        doc.add_heading(section_titles.get(key, key.title()), level=1)
        for line in text.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith("---"):
                doc.add_paragraph("—" * 30)
            else:
                doc.add_paragraph(stripped)

        doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_title = "".join(c if c.isalnum() or c in "-_ " else "_" for c in story.title)[:60]
    filename = f"{safe_title}_story_bible_v{bible.version}.docx"

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
