import io
import re
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session

from database import get_db
from models import Story, Chapter, User as UserModel
from schemas import ExportRequest
from routers.auth import get_current_user

router = APIRouter(tags=["export"])

# Sanitise a story title for use as a filename — strips filesystem-unsafe chars
_UNSAFE_CHARS = re.compile(r'[\\/*?:"<>|]')


def _safe_filename(title: str, ext: str) -> str:
    sanitised = _UNSAFE_CHARS.sub("", title).strip()[:100] or "manuscript"
    return f"{sanitised}.{ext}"


@router.post("/")
async def export_manuscript(
    data: ExportRequest,
    current_user: UserModel = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    story = db.query(Story).filter(
        Story.story_id == data.story_id,
        Story.user_id == current_user.user_id,
    ).first()
    if not story:
        raise HTTPException(status_code=404, detail="Story not found")

    chapters = (
        db.query(Chapter)
        .filter(Chapter.story_id == data.story_id)
        .order_by(Chapter.chapter_number)
        .all()
    )

    author_name: str = current_user.username
    font_pt = max(8, min(int(data.font_size or 12), 24))

    if data.format == "docx":
        return _export_docx(story, chapters, data, author_name, font_pt)
    elif data.format == "pdf":
        return _export_pdf(story, chapters, data, author_name, font_pt)
    else:
        raise HTTPException(status_code=400, detail="Format must be 'docx' or 'pdf'")


# ── DOCX export ────────────────────────────────────────────────────────────────

def _export_docx(story, chapters, data, author_name: str, font_pt: int):
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Document metadata
    doc.core_properties.author = author_name
    doc.core_properties.title = story.title

    # Standard manuscript margins: 1.25 in left/right, 1.0 in top/bottom
    for section in doc.sections:
        section.page_width    = Inches(8.5)
        section.page_height   = Inches(11)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    # ── Title page ─────────────────────────────────────────────────────────────
    # Push content toward vertical centre with leading whitespace
    for _ in range(8):
        blank = doc.add_paragraph()
        blank.space_after = Pt(0)

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(story.title)
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(28)
    title_run.bold = True

    doc.add_paragraph()

    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_run = author_para.add_run(f"by {author_name}")
    author_run.font.name = "Times New Roman"
    author_run.font.size = Pt(16)

    if story.description:
        doc.add_paragraph()
        desc_para = doc.add_paragraph()
        desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        desc_run = desc_para.add_run(story.description)
        desc_run.font.name = "Times New Roman"
        desc_run.font.size = Pt(12)
        desc_run.italic = True

    doc.add_page_break()

    # ── Chapters ───────────────────────────────────────────────────────────────
    for idx, ch in enumerate(chapters):
        heading_text = (
            f"Chapter {ch.chapter_number}: {ch.title}"
            if data.include_chapter_numbers
            else ch.title
        )

        heading_para = doc.add_heading(heading_text, level=1)
        heading_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in heading_para.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(18)

        doc.add_paragraph()

        content = (ch.content or "").strip()
        if content:
            for para_text in content.split("\n\n"):
                para_text = para_text.strip()
                if not para_text:
                    continue
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                run = p.add_run(para_text.replace("\n", " "))
                run.font.name = "Times New Roman"
                run.font.size = Pt(font_pt)

        if idx < len(chapters) - 1:
            doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = _safe_filename(story.title, "docx")
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ── PDF export ─────────────────────────────────────────────────────────────────

def _export_pdf(story, chapters, data, author_name: str, font_pt: int):
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak

    buf = io.BytesIO()

    def _page_number(canvas, doc):
        # No number on the title page (page 1)
        if canvas.getPageNumber() == 1:
            return
        canvas.saveState()
        canvas.setFont("Times-Roman", 9)
        canvas.drawCentredString(4.25 * inch, 0.5 * inch, str(canvas.getPageNumber()))
        canvas.restoreState()

    pdf_doc = SimpleDocTemplate(
        buf,
        pagesize=LETTER,
        rightMargin=1.25 * inch,
        leftMargin=1.25 * inch,
        topMargin=1.0 * inch,
        bottomMargin=1.0 * inch,
        title=story.title,
        author=author_name,
    )

    title_style = ParagraphStyle(
        "StoryTitle",
        fontName="Times-Bold",
        fontSize=28,
        leading=36,
        alignment=1,
        spaceAfter=18,
    )
    author_style = ParagraphStyle(
        "StoryAuthor",
        fontName="Times-Italic",
        fontSize=16,
        leading=22,
        alignment=1,
        spaceAfter=12,
    )
    desc_style = ParagraphStyle(
        "StoryDesc",
        fontName="Times-Italic",
        fontSize=12,
        leading=18,
        alignment=1,
        spaceAfter=12,
    )
    heading_style = ParagraphStyle(
        "ChapterHeading",
        fontName="Times-Bold",
        fontSize=18,
        leading=26,
        alignment=1,
        spaceBefore=0,
        spaceAfter=14,
    )
    body_style = ParagraphStyle(
        "StoryBody",
        fontName="Times-Roman",
        fontSize=font_pt,
        leading=font_pt * 1.6,
        firstLineIndent=0.3 * inch,
        spaceAfter=6,
        alignment=4,  # justified
    )

    def _xml_escape(text: str) -> str:
        return (
            text
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        )

    elements = []

    # ── Title page ─────────────────────────────────────────────────────────────
    elements.append(Spacer(1, 2.5 * inch))
    elements.append(Paragraph(_xml_escape(story.title), title_style))
    elements.append(Spacer(1, 0.5 * inch))
    elements.append(Paragraph(f"by {_xml_escape(author_name)}", author_style))
    if story.description:
        elements.append(Spacer(1, 0.5 * inch))
        elements.append(Paragraph(_xml_escape(story.description), desc_style))
    elements.append(PageBreak())

    # ── Chapters ───────────────────────────────────────────────────────────────
    for idx, ch in enumerate(chapters):
        heading_text = (
            f"Chapter {ch.chapter_number}: {ch.title}"
            if data.include_chapter_numbers
            else ch.title
        )
        elements.append(Paragraph(_xml_escape(heading_text), heading_style))
        elements.append(Spacer(1, 0.2 * inch))

        content = (ch.content or "").strip()
        if content:
            for para_text in content.split("\n\n"):
                para_text = para_text.strip()
                if not para_text:
                    continue
                safe = _xml_escape(para_text).replace("\n", " ")
                elements.append(Paragraph(safe, body_style))

        if idx < len(chapters) - 1:
            elements.append(PageBreak())

    pdf_doc.build(elements, onFirstPage=_page_number, onLaterPages=_page_number)
    buf.seek(0)

    filename = _safe_filename(story.title, "pdf")
    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
